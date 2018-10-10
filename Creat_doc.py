#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Inches

from main import *
import matplotlib.pyplot as plt
import numpy as np


#document.add_page_break()
speed=[]
Trac_Force=[]

for i in range(1,len(Trac_Chara_speed_List)):
	speed.append(Trac_Chara_speed_List[i][0])
	Trac_Force.append(Trac_Chara_speed_List[i][1])


#plt.figure(figsize=(500, 500))
plt.plot(speed,Trac_Force)

plt.xlim(0, 85)
plt.ylim(0, 450)
my_x_ticks = np.arange(0, 85, 5)
my_y_ticks = np.arange(0, 450, 50)

plt.xticks(my_x_ticks)
plt.yticks(my_y_ticks)

plt.savefig('test.jpg')

document=Document()

document.add_heading('HEAD',0)

document.add_picture('test.jpg',width=Inches(6))
#plt.show()



document.save('test.docx')

